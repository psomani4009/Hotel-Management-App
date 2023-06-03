from django.shortcuts import render
from django.http import HttpResponse, HttpResponseRedirect
from django.template import loader
import pandas as pd
import docx, datetime, os
from pytz import timezone
# Create your views here.

def home(request):
    template = loader.get_template('index.html')
    try:
        pd.read_csv('data.csv')
    except:
        f = open('data.csv', 'w')
        f.write('SlNo,roomno,name,address,idType,idNum,pov,phone,checkin,checkout')
        f.close()
    f = open("data.csv")
    r = f.read().split('\n')[1:]
    stay = []
    for i in r:
        z = i.strip().split(',')
        if z[-1] == '':
            stay.append(z[1])
    context = {'booked': stay}
    return HttpResponse(template.render(context))

def dorm(request):
    try:
        pd.read_csv('data.csv')
    except:
        f = open('data.csv', 'w')
        f.write('SlNo,roomno,name,address,idType,idNum,pov,phone,checkin,checkout')
        f.close()
    f = open("data.csv")
    r = f.read().split('\n')[1:]
    stay = []
    for i in r:
        z = i.strip().split(',')
        if z[-1] == '':
            stay.append(z[1])
    context = {'booked': stay}
    template = loader.get_template("dorm.html")
    return HttpResponse(template.render(context, request))

def assign_room(request, roomno):
    if request.method == 'POST':
        f = open('data.csv', 'a')
        f.write('\n'+','.join([str(i) for i in [request.POST.get('SlNo'), request.POST.get('roomno'),request.POST.get('name'), request.POST.get('address'), request.POST.get('idType'), request.POST.get('idNum'),request.POST.get('pov'), request.POST.get('phone'), request.POST.get('checkin'), request.POST.get('checkout')]]))
        f.close()
        return HttpResponseRedirect('/')
    template = loader.get_template('form.html')
    context = {
        'room_no': roomno,
        'dt': datetime.datetime.now(timezone('Asia/Kolkata')).strftime('%Y-%m-%dT%H:%M')
    }
    return HttpResponse(template.render(context, request))

def checkout_room(request, roomno):
    if request.method == "POST":
        f = open("data.csv", "r")
        r = f.read()
        h = [r.split('\n')[0]]
        d = r.split('\n')[1:]
        d[int(request.POST.get('index'))] = ','.join([str(i) for i in [request.POST.get('SlNo'), request.POST.get('roomno'),request.POST.get('name'), request.POST.get('address'), request.POST.get('idType'), request.POST.get('idNum'),request.POST.get('pov'), request.POST.get('phone'), request.POST.get('checkin'), request.POST.get('checkout')]])
        f.close()
        f = open('data.csv', "w")
        s = '\n'.join(h+d)
        f.write(s)
        f.close()
        return HttpResponseRedirect('/')
    template = loader.get_template('checkout.html')
    f = open('data.csv', 'r')
    r = f.read()
    d = r.split('\n')[1:]
    index = -1
    for i in range(len(d)):
        data = d[i].strip().split(",")
        if len(data)>1 and data[1] == str(roomno) and data[-1] == '':
            index = i
            break
    context = {
        "index": index,
        "roomno": roomno,
        "data": data
    }
    return HttpResponse(template.render(context, request))

def gen(request):
    if request.method == 'POST':
        doc = docx.Document()
        p = doc.add_paragraph('Mobile No.: 00000 00000')
        p.alignment = 2
        head = doc.add_heading("Hotel Xyz Inn", 0)
        head.alignment = 1
        doc.add_paragraph('Date: ')
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        row = table.rows[0].cells
        row[0].text = "Sl. No."
        row[1].text = "Name"
        row[2].text = "Address"
        row[3].text = "ID Type"
        row[4].text = "ID"
        row[5].text = "POV"
        row[6].text = "Mobile No."
        f = open("data.csv", 'r')
        r = f.read().strip().split('\n')[1:]
        c = 1
        if request.POST.get('checkin') == '':
            for i in r:
                z = i.split(',')
                if (z[-1] == ''):
                    row = table.add_row().cells
                    row[0].text = str(c)
                    row[1].text = z[2]
                    row[2].text = z[3]
                    row[3].text = z[4]
                    row[4].text = z[5]
                    row[5].text = z[6]
                    row[6].text = z[7]
                    c+=1
        elif request.POST.get('checkout') == '':
            dt = datetime.datetime.strptime(request.POST.get('checkin').replace('T', ' '), '%Y-%m-%d %H:%M')
            for i in r:
                z = i.split(',')
                checkin = datetime.datetime.strptime(z[-2].replace('T', ' '), '%Y-%m-%d %H:%M')
                if dt.timestamp()<checkin.timestamp() and z[-1] == '':
                    row = table.add_row().cells
                    row[0].text = str(c)
                    row[1].text = z[2]
                    row[2].text = z[3]
                    row[3].text = z[4]
                    row[4].text = z[5]
                    row[5].text = z[6]
                    row[6].text = z[7]
                    c+=1
        else:
            pci = datetime.datetime.strptime(request.POST.get('checkin').replace('T', ' '), '%Y-%m-%d %H:%M')
            pco = datetime.datetime.strptime(request.POST.get('checkout').replace('T', ' '), '%Y-%m-%d %H:%M')
            for i in r:
                z = i.split(',')
                checkin = datetime.datetime.strptime(z[-2].replace('T', ' '), '%Y-%m-%d %H:%M')
                checkout = datetime.datetime.strptime(z[-1].replace('T', ' '), '%Y-%m-%d %H:%M') if z[-1] != '' else pci
                if pci.timestamp()<checkin.timestamp() and checkout.timestamp()<pco.timestamp():
                    row = table.add_row().cells
                    row[0].text = str(c)
                    row[1].text = z[2]
                    row[2].text = z[3]
                    row[3].text = z[4]
                    row[4].text = z[5]
                    row[5].text = z[6]
                    row[6].text = z[7]
                    c+=1
        if os.name == 'posix':
            desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') 
        else:
            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
        doc.save(os.path.join(desktop, f'Hotel Xyz Inn {request.POST.get("filename")}.docx'))
        os.startfile(os.path.join(desktop, f'Hotel Xyz Inn {request.POST.get("filename")}.docx'))
        return HttpResponseRedirect('/')
    template = loader.get_template('report.html')
    return HttpResponse(template.render({}, request))